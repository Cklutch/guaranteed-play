from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("outputs/guaranteed_play_backend_model_update.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(col)


def apply_run_style(run, bold=False, color=INK, size=11):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        apply_run_style(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        apply_run_style(r2)
    else:
        run = p.add_run(text)
        apply_run_style(run)
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE
    return p


def add_callout(doc, label, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r1 = p.add_run(label + " ")
    apply_run_style(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    apply_run_style(r2)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        set_cell_margins(hdr[idx])
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        apply_run_style(r, bold=True, color=DARK_BLUE)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            apply_run_style(r, size=10)

    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Guaranteed Play Backend Model Update")
    apply_run_style(r, bold=True, color=RGBColor(11, 37, 69), size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(
        f"Current-state briefing for the project owner | {date(2026, 6, 25).strftime('%B %-d, %Y') if False else 'June 25, 2026'}"
    )
    apply_run_style(r, color=MUTED, size=10)

    add_callout(
        doc,
        "Bottom line:",
        "Guaranteed Play currently functions as an interpretable draft-decision engine, not a frozen predictive model. It is ahead of a normal cheat sheet because it adapts to roster construction, ADP, tiers, availability, and simulated championship impact. Statistically, the next maturity step is validation: backtests, calibration, benchmark comparisons, and frozen train/test definitions.",
    )

    add_heading(doc, "1. Where We Are", 1)
    add_bullet(doc, "Backend focus is now the draft model, scoring logic, simulation logic, and research validation path.")
    add_bullet(doc, "The current app-facing model is the master recommendation engine in draftkit/draft_analysis.py, with a consensus layer on top.")
    add_bullet(doc, "The model is interpretable and modular: each major recommendation can be explained by component scores rather than hidden black-box behavior.")
    add_bullet(doc, "WR and RB underpriced research is promising but remains outside the production draft model.")
    add_bullet(doc, "The most important backend risk is not lack of signal. It is insufficient frozen validation against historical outcomes and industry baselines.")

    add_heading(doc, "2. How the Current Model Functions", 1)
    add_para(
        doc,
        "The model starts with the available player pool, removes drafted players, detects the user's roster needs, then scores remaining QB/RB/WR/TE players. Its first job is not to predict a player's exact season. Its practical job is to answer: given the current pick, roster, league settings, and market price, who is the best draft decision right now?",
    )
    add_para(
        doc,
        "The core model score is a weighted blend of five normalized inputs:",
    )
    add_table(
        doc,
        ["Component", "Current Weight", "Plain-English Meaning"],
        [
            ["Projection value", "30%", "How much projected value the player offers after adjusting for position and replacement level."],
            ["Position need", "25%", "How badly the current roster needs the player's position."],
            ["ADP value", "15%", "Whether the player is cheaper than his projected/ranked value suggests."],
            ["Tier urgency", "15%", "Whether waiting risks falling into a worse tier at the position."],
            ["Team fit", "15%", "Whether the player's risk/upside profile balances the roster already drafted."],
        ],
        [1.65, 1.05, 3.8],
    )
    add_para(
        doc,
        "After that base blend, the backend applies guardrails. Construction pressure boosts positions that are becoming urgent. Signal trust can cap or reduce a score when ADP, projection, or market inputs look inconsistent. Single-QB leagues cap QB values so raw QB points do not overwhelm RB/WR economics.",
    )
    add_para(
        doc,
        "Draft Mode then builds a consensus score. This blends the base model score with value, roster fit, future-pick impact, championship equity, safety/risk, and survival probability. This consensus layer is the current top-level decision layer.",
    )

    add_heading(doc, "3. Simulation Layers", 1)
    add_bullet(doc, "Availability simulation estimates whether a player will survive until the user's next pick using ADP-weighted Monte Carlo paths.")
    add_bullet(doc, "Future-impact simulation asks what the board may look like after taking a candidate now and which top targets are likely to be lost or survive.")
    add_bullet(doc, "Championship equity V2 estimates directional roster impact using floor, median, ceiling, injury risk, role uncertainty, age risk, roster shape, and league-winning upside.")
    add_bullet(doc, "These simulations make the model more draft-aware than a static ranking list, but their outputs should be treated as directional, not literal probabilities.")

    add_heading(doc, "4. Comparison to Fantasy Football Industry Standards", 1)
    add_para(
        doc,
        "The fantasy football industry usually starts from projected points, expert consensus rankings, ADP, tiers, injury risk, positional scarcity, and draft strategy rules. More advanced products layer in value-based drafting, player archetypes, risk, opportunity, and mock-draft behavior.",
    )
    add_table(
        doc,
        ["Industry Practice", "Guaranteed Play Today", "Assessment"],
        [
            ["Projected points / rankings", "Uses projections and projection rank as a core input.", "Meets standard."],
            ["Value-based drafting", "Uses position-adjusted value over replacement and ADP discounts.", "Meets and extends standard."],
            ["ADP awareness", "Compares projection rank to ADP rank and models survival to the next pick.", "Above basic standard."],
            ["Tier drafting", "Calculates tier urgency, tier drop-off, and desperation targets.", "Meets standard."],
            ["Roster construction", "Dynamically adjusts by current roster, position needs, and RB/WR imbalance.", "Above static cheat-sheet standard."],
            ["Risk/upside", "Uses injury risk, durability, boom/bust/stability, archetype, and safety scoring.", "Meets standard, but source quality matters."],
            ["Simulation", "Uses Monte Carlo availability/future-board/championship equity layers.", "Advanced for a personal draft tool."],
            ["Historical validation", "Research exists, but production model is not fully benchmark-frozen.", "Below mature analytics standard."],
        ],
        [1.8, 2.65, 2.05],
    )
    add_para(
        doc,
        "Compared with a typical fantasy platform, the backend is strongest in dynamic draft context. Compared with a mature paid analytics model, it still needs stronger validation, calibration, and documented benchmark performance.",
    )

    add_heading(doc, "5. Comparison to Statistical Modeling Standards", 1)
    add_para(
        doc,
        "From a statistics perspective, the current backend is best described as an interpretable scoring model plus simulation layer. It is not yet a fully validated supervised predictive model for season outcomes.",
    )
    add_bullet(doc, "Strength: features are explainable and aligned with draft decision-making: projection, market price, scarcity, roster need, risk, and simulation outcomes.")
    add_bullet(doc, "Strength: the RB research result uses a regularized logistic regression framing, which is statistically sensible for underpriced-hit probability.")
    add_bullet(doc, "Strength: hit rate, base rate, lift, ROC AUC, and model family are already being tracked in research notes.")
    add_bullet(doc, "Gap: the app-facing model lacks a frozen historical backtest with train/test separation by season.")
    add_bullet(doc, "Gap: the model score is not calibrated to a real probability, so a 90 score should not be read as a 90% chance of success.")
    add_bullet(doc, "Gap: there is not yet a benchmark table against simple baselines such as ADP-only, projection-only, ECR-only, and value-over-replacement-only.")
    add_bullet(doc, "Gap: uncertainty intervals and sensitivity checks are not yet shown for simulation outputs.")
    add_para(
        doc,
        "Industry-grade statistical practice would ask the model to prove that it beats simpler baselines out of sample, that probability outputs are calibrated, and that performance remains stable across seasons, positions, draft ranges, and scoring formats.",
    )

    add_heading(doc, "6. Trusted vs Experimental", 1)
    add_table(
        doc,
        ["Area", "Current Status", "Use Right Now"],
        [
            ["Master recommendation score", "Trusted directionally", "Use as the core ranking signal."],
            ["Consensus score", "Trusted directionally", "Use as the Draft Mode decision layer."],
            ["ADP value", "Trusted with guardrails", "Use when ADP/projection inputs pass trust checks."],
            ["Availability probability", "Directional", "Use for wait/take guidance, not exact odds."],
            ["Championship equity V2", "Directional", "Use for upside/portfolio context, not literal title odds."],
            ["WR underpriced models", "Research-only", "Do not wire into Draft Mode yet."],
            ["RB underpriced model", "Promising research-only", "Benchmark/freeze audit before app use."],
            ["Score V4", "Paused", "Do not use until compression issue is resolved."],
        ],
        [1.9, 1.85, 2.75],
    )

    add_heading(doc, "7. Recommended Backend-Only Roadmap", 1)
    add_para(doc, "The next work should focus on model proof, not UI cleanup.")
    add_bullet(doc, "Freeze definitions: define target outcomes for WR2, WR1 ceiling, RB2, league-winning picks, and bust outcomes.")
    add_bullet(doc, "Build baselines: compare every research model against ADP-only, projection-only, ECR-only, and simple value-over-replacement rankings.")
    add_bullet(doc, "Backtest by season: train on older seasons and test on held-out future seasons to reduce leakage.")
    add_bullet(doc, "Calibrate probabilities: if a model says 40%, players in that bucket should hit about 40% over time.")
    add_bullet(doc, "Separate ranking scores from probabilities: keep model score for ordering, but reserve probability language for calibrated models.")
    add_bullet(doc, "Audit feature leakage: verify that no future-season outcome data sneaks into draft-time features.")
    add_bullet(doc, "Promote only frozen signals: wire research into Draft Mode only after benchmark, calibration, and documentation pass.")

    add_heading(doc, "8. Owner Takeaway", 1)
    add_callout(
        doc,
        "Current position:",
        "Guaranteed Play has the bones of a strong backend draft intelligence system. Its edge is contextual decision-making, not just player ranking. The next serious step is to turn promising research into validated, frozen model components that can survive comparison against simple industry baselines.",
        fill="F4F6F9",
    )

    add_heading(doc, "Sources Used for Industry and Statistics Comparison", 1)
    sources = [
        "Fantasy football overview and draft concepts, including ADP and value-based drafting: https://en.wikipedia.org/wiki/Fantasy_football_(gridiron)",
        "Scikit-learn model evaluation guidance, including proper metrics, Brier score, log loss, ROC AUC, and baseline estimators: https://scikit-learn.org/stable/modules/model_evaluation.html",
        "Scikit-learn probability calibration guidance: https://scikit-learn.org/stable/modules/calibration.html",
        "Scikit-learn cross-validation guidance: https://scikit-learn.org/stable/modules/cross_validation.html",
        "Dynamic fantasy roster/value modeling research: https://arxiv.org/abs/2409.09884",
        "Fantasy sports optimization research using predictive modeling, robust optimization, and Monte Carlo simulation: https://arxiv.org/abs/2505.02170",
        "Model evaluation caution around AUC and the need for threshold/base-rate context: https://arxiv.org/abs/2305.18159",
    ]
    for source in sources:
        add_bullet(doc, source)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Guaranteed Play backend model update")
    apply_run_style(r, color=MUTED, size=9)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH.resolve())
